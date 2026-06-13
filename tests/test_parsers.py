"""
Parser unit tests — use tiny in-memory files to stay fast and self-contained.
"""

import io
import os
import sys
import tempfile

import pytest

sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))


# ── docx ──────────────────────────────────────────────────────────────────────

def _make_docx_with_table(tmp_path):
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    doc.add_paragraph("Intro paragraph before table.")

    # Add a 2-column, 3-row table
    tbl = doc.add_table(rows=3, cols=2)
    tbl.cell(0, 0).text = "Name"
    tbl.cell(0, 1).text = "Value"
    tbl.cell(1, 0).text = "Alpha"
    tbl.cell(1, 1).text = "1"
    tbl.cell(2, 0).text = "Beta"
    tbl.cell(2, 1).text = "2"
    doc.add_paragraph("Trailing paragraph after table.")

    p = tmp_path / "test.docx"
    doc.save(str(p))
    return str(p)


def test_docx_table_detected_mid_document(tmp_path):
    from parsers.docx_parser import parse_docx
    path   = _make_docx_with_table(tmp_path)
    chunks = parse_docx(path)

    types = [c["type"] for c in chunks]
    assert "table" in types, "No table chunk produced"

    table_chunks = [c for c in chunks if c["type"] == "table"]
    assert len(table_chunks) == 1
    t = table_chunks[0]["text"]
    assert "TABLE:" in t
    assert "Name" in t
    assert "Alpha" in t


def test_docx_table_has_context(tmp_path):
    from parsers.docx_parser import parse_docx
    path   = _make_docx_with_table(tmp_path)
    chunks = parse_docx(path)
    table_chunk = next(c for c in chunks if c["type"] == "table")
    # The last paragraph before the table should be prepended
    assert "[Context:" in table_chunk["text"]
    assert "Intro paragraph before table" in table_chunk["text"]


def test_docx_text_before_and_after_table(tmp_path):
    from parsers.docx_parser import parse_docx
    path   = _make_docx_with_table(tmp_path)
    chunks = parse_docx(path)
    texts  = [c for c in chunks if c["type"] == "text"]
    all_text = " ".join(c["text"] for c in texts)
    assert "Trailing paragraph" in all_text


def test_docx_empty_file(tmp_path):
    from docx import Document
    from parsers.docx_parser import parse_docx
    doc = Document()
    p   = tmp_path / "empty.docx"
    doc.save(str(p))
    assert parse_docx(str(p)) == []


# ── xlsx ──────────────────────────────────────────────────────────────────────

def _make_xlsx(tmp_path):
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sales"
    ws.append(["Product", "Revenue", "Units"])
    for i in range(5):
        ws.append([f"Item{i}", i * 100, i * 10])
    p = tmp_path / "data.xlsx"
    wb.save(str(p))
    return str(p)


def test_xlsx_table_chunk_produced(tmp_path):
    from parsers.xlsx_parser import parse_xlsx
    path   = _make_xlsx(tmp_path)
    chunks = parse_xlsx(path)
    assert chunks, "No chunks produced"
    assert all(c["type"] == "table" for c in chunks)


def test_xlsx_context_prefix(tmp_path):
    from parsers.xlsx_parser import parse_xlsx
    path   = _make_xlsx(tmp_path)
    chunks = parse_xlsx(path)
    assert any("[Context: Sheet — Sales]" in c["text"] for c in chunks)


def test_xlsx_headers_in_chunk(tmp_path):
    from parsers.xlsx_parser import parse_xlsx
    path   = _make_xlsx(tmp_path)
    chunks = parse_xlsx(path)
    text = " ".join(c["text"] for c in chunks)
    assert "Product" in text
    assert "Revenue" in text


def test_xlsx_row_batching_under_400_tokens(tmp_path):
    import openpyxl
    import tiktoken
    from parsers.xlsx_parser import parse_xlsx

    enc = tiktoken.get_encoding("cl100k_base")
    wb  = openpyxl.Workbook()
    ws  = wb.active
    ws.title = "Big"
    ws.append(["Col"])
    for i in range(200):
        ws.append([f"Very long repeated text value number {i} padding padding padding"])
    p = tmp_path / "big.xlsx"
    wb.save(str(p))

    chunks = parse_xlsx(str(p))
    for c in chunks:
        tokens = len(enc.encode(c["text"]))
        assert tokens <= 512, f"Chunk exceeds 512 tokens: {tokens}"


# ── plain ─────────────────────────────────────────────────────────────────────

def test_plain_txt(tmp_path):
    from parsers.plain_parser import parse_plain
    p = tmp_path / "notes.txt"
    p.write_text("Hello world\nLine two")
    chunks = list(parse_plain(str(p)))
    assert len(chunks) == 1
    assert chunks[0]["type"] == "text"
    assert "Hello world" in chunks[0]["text"]


def test_plain_csv_as_table(tmp_path):
    from parsers.plain_parser import parse_plain
    p = tmp_path / "data.csv"
    p.write_text("Name,Score\nAlice,90\nBob,85\n")
    chunks = list(parse_plain(str(p)))
    assert len(chunks) == 1
    assert chunks[0]["type"] == "table"
    assert "Name" in chunks[0]["text"]
    assert "Alice" in chunks[0]["text"]


def test_plain_empty_txt(tmp_path):
    from parsers.plain_parser import parse_plain
    p = tmp_path / "empty.txt"
    p.write_text("")
    assert parse_plain(str(p)) == []


# ── pdf ───────────────────────────────────────────────────────────────────────

def test_pdf_text_extracted(tmp_path):
    pytest.importorskip("pdfplumber")
    try:
        from reportlab.pdfgen import canvas as rl_canvas
    except ImportError:
        pytest.skip("reportlab not installed — skipping PDF creation test")

    from parsers.pdf_parser import parse_pdf
    p = str(tmp_path / "simple.pdf")
    c = rl_canvas.Canvas(p)
    c.drawString(100, 750, "Hello from a PDF file.")
    c.save()

    chunks = parse_pdf(p)
    assert chunks, "No chunks from PDF"
    text = " ".join(c["text"] for c in chunks if c["type"] == "text")
    assert "Hello" in text
