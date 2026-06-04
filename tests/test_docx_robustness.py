from docx import Document
from parsers.docx_parser import parse_docx


def test_body_text_parses(tmp_path):
    d = Document(); d.add_paragraph("Body content here.")
    f = tmp_path / "n.docx"; d.save(str(f))
    assert len(parse_docx(str(f))) >= 1


def test_header_only_is_recovered(tmp_path):
    d = Document(); d.sections[0].header.paragraphs[0].text = "CONFIDENTIAL header text"
    f = tmp_path / "h.docx"; d.save(str(f))
    assert len(parse_docx(str(f))) >= 1


def test_footer_only_is_recovered(tmp_path):
    d = Document(); d.sections[0].footer.paragraphs[0].text = "Footer legal text"
    f = tmp_path / "ft.docx"; d.save(str(f))
    assert len(parse_docx(str(f))) >= 1


def test_truly_empty_stays_empty(tmp_path):
    d = Document(); f = tmp_path / "e.docx"; d.save(str(f))
    assert len(parse_docx(str(f))) == 0
