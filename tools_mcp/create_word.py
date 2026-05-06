"""
create_word_document tool — build a real .docx from structured content
using python-docx, upload to OneDrive.
"""

import io
from typing import Literal

from fastapi import APIRouter, Depends, Request
from pydantic import BaseModel, Field
from slowapi import Limiter
from slowapi.util import get_remote_address

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from tools_mcp.auth import require_bearer
from tools_mcp._onedrive_upload import upload_bytes

router  = APIRouter()
limiter = Limiter(key_func=get_remote_address)

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class WordTable(BaseModel):
    headers: list[str] = Field(default_factory=list)
    rows:    list[list[str]] = Field(default_factory=list)
    caption: str | None = None


class WordSection(BaseModel):
    heading:    str | None = None
    heading_level: int = 1  # 1..4
    paragraphs: list[str] = Field(default_factory=list)
    bullets:    list[str] = Field(default_factory=list)
    table:      WordTable | None = None


class CreateWordRequest(BaseModel):
    filename:    str
    folder_path: str
    title:       str | None = None
    subtitle:    str | None = None
    sections:    list[WordSection] = Field(default_factory=list)


class CreateWordResponse(BaseModel):
    success:      bool
    onedrive_url: str | None = None
    file_id:      str | None = None
    error:        str | None = None


def _build_docx(req: CreateWordRequest) -> bytes:
    doc = Document()

    # Default body font size
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    if req.title:
        t = doc.add_heading(req.title, level=0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if req.subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(req.subtitle)
        run.italic = True
        run.font.size = Pt(12)

    for sec in req.sections:
        if sec.heading:
            level = max(1, min(sec.heading_level, 4))
            doc.add_heading(sec.heading, level=level)

        for para in sec.paragraphs:
            doc.add_paragraph(para)

        for bullet in sec.bullets:
            doc.add_paragraph(bullet, style="List Bullet")

        if sec.table and (sec.table.headers or sec.table.rows):
            headers = sec.table.headers
            rows    = sec.table.rows
            ncols   = max(len(headers), max((len(r) for r in rows), default=0))
            if ncols > 0:
                if sec.table.caption:
                    cap = doc.add_paragraph()
                    cap_run = cap.add_run(sec.table.caption)
                    cap_run.bold = True

                table = doc.add_table(rows=1 if headers else 0, cols=ncols)
                table.style = "Light Grid Accent 1"

                if headers:
                    hdr_cells = table.rows[0].cells
                    for i in range(ncols):
                        text = headers[i] if i < len(headers) else ""
                        hdr_cells[i].text = text
                        for run in hdr_cells[i].paragraphs[0].runs:
                            run.bold = True

                for row in rows:
                    cells = table.add_row().cells
                    for i in range(ncols):
                        cells[i].text = row[i] if i < len(row) else ""

                doc.add_paragraph()  # spacer after table

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.post("/tools/create_word_document", response_model=CreateWordResponse)
@limiter.limit("10/minute")
def create_word_document(
    request: Request,
    req: CreateWordRequest,
    _: None = Depends(require_bearer),
) -> CreateWordResponse:
    try:
        filename = req.filename if req.filename.lower().endswith(".docx") else f"{req.filename}.docx"
        raw = _build_docx(req)
        data = upload_bytes(raw, filename, req.folder_path, _DOCX_MIME)
        return CreateWordResponse(
            success      = True,
            onedrive_url = data.get("webUrl"),
            file_id      = data.get("id"),
        )
    except Exception as e:
        return CreateWordResponse(success=False, error=str(e))
